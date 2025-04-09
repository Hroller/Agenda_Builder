from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime

# Define the agenda categories
CATEGORIES = [
    "Unapproved minutes from the last meeting",
    "Petitions",
    "Reports from Committees",
    "Financial Report",
    "Sickness & Distress",
    "Communication",
    "Unfinished Business",
    "New Business"
]

# Set document margins
def set_margins(section):
    section.left_margin = Inches(0.2)
    section.right_margin = Inches(0.2)

# Ask for date(s) of unapproved minutes; if the input is 0, O, none, or empty, return ["None"].
def get_date_input(prompt):
    while True:
        date_input = input(prompt)
        if date_input.strip().lower() in ("0", "o", "none", ""):
            return ["None"]
        dates = [d.strip() for d in date_input.split(',')]
        try:
            formatted = [datetime.strptime(d, '%Y-%m-%d').strftime('%B %d, %Y') for d in dates]
            return formatted
        except ValueError:
            print("Please enter the date(s) in YYYY-MM-DD format, or enter 0 (or O/none) if there are none.")

# Ask for agenda items per category
def collect_agenda_items():
    agenda_items = {}
    for category in CATEGORIES:
        if category == "Unapproved minutes from the last meeting":
            dates = get_date_input("Enter the date(s) of the unapproved minutes (comma-separated, YYYY-MM-DD, or 0/O/none if none): ")
            agenda_items[category] = dates
        else:
            items = input(f"Enter items for '{category}' (comma-separated): ").split(',')
            agenda_items[category] = [item.strip() for item in items if item.strip()]
    return agenda_items

# Create .docx agenda
def create_agenda_docx(agenda_items, filename, meeting_date, group_name):
    doc = Document()
    set_margins(doc.sections[0])

    heading = doc.add_paragraph(f"{group_name} Agenda for {meeting_date}")
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = heading.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    p = doc.add_paragraph("Greet all MEC/VEC/ECs")
    p.runs[0].font.name = 'Times New Roman'
    p.runs[0].font.size = Pt(9)
    for _ in range(3):
        line = doc.add_paragraph("    ____________________________________________")
        line.runs[0].font.name = 'Times New Roman'
        line.runs[0].font.size = Pt(9)

    for index, category in enumerate(CATEGORIES, 1):
        p = doc.add_paragraph(f"{index}. {category}")
        p.runs[0].font.name = 'Times New Roman'
        p.runs[0].font.size = Pt(9)
        for item in agenda_items[category]:
            line = doc.add_paragraph(f"    - {item}")
            line.runs[0].font.name = 'Times New Roman'
            line.runs[0].font.size = Pt(9)

    doc.save(filename)
    print(f"Agenda created: {filename}")

# Create .docx minutes
def create_minutes_docx(agenda_items, filename, meeting_date, group_name):
    doc = Document()
    set_margins(doc.sections[0])

    heading = doc.add_paragraph(f"Meeting minutes for {group_name} for {meeting_date}")
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = heading.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    p = doc.add_paragraph("Opening Time: __________________")
    p.runs[0].font.name = 'Times New Roman'
    p.runs[0].font.size = Pt(9)

    p = doc.add_paragraph("Number of Members Present: __________    Number of Visitors Present: __________    Total: _____")
    p.runs[0].font.name = 'Times New Roman'
    p.runs[0].font.size = Pt(9)

    p = doc.add_paragraph("MEC/VEC/ECs Present:")
    p.runs[0].font.name = 'Times New Roman'
    p.runs[0].font.size = Pt(9)
    for _ in range(3):
        line = doc.add_paragraph("    ____________________________________________")
        line.runs[0].font.name = 'Times New Roman'
        line.runs[0].font.size = Pt(9)

    for index, category in enumerate(CATEGORIES, 1):
        p = doc.add_paragraph(f"{index}. {category}")
        p.runs[0].font.name = 'Times New Roman'
        p.runs[0].font.size = Pt(9)
        for item in agenda_items[category]:
            line = doc.add_paragraph(f"    - {item}")
            line.runs[0].font.name = 'Times New Roman'
            line.runs[0].font.size = Pt(9)

    p = doc.add_paragraph("Closing Time: __________________")
    p.runs[0].font.name = 'Times New Roman'
    p.runs[0].font.size = Pt(9)

    doc.save(filename)
    print(f"Minutes created: {filename}")

# Create plain text agenda
def create_agenda_txt(agenda_items, filename, meeting_date, group_name):
    with open(filename, 'w') as f:
        f.write(f"{group_name} Agenda for {meeting_date}\n\n")
        f.write("Greet all MEC/VEC/ECs\n")
        f.writelines(["    ____________________________________________\n" for _ in range(3)])
        for index, category in enumerate(CATEGORIES, 1):
            f.write(f"{index}. {category}\n")
            for item in agenda_items[category]:
                f.write(f"    - {item}\n")
            f.write("\n")
    print(f"Agenda created: {filename}")

# Create plain text minutes
def create_minutes_txt(agenda_items, filename, meeting_date, group_name):
    with open(filename, 'w') as f:
        f.write(f"Meeting minutes for {group_name} for {meeting_date}\n\n")
        f.write("Opening Time: __________________\n")
        f.write("Number of Members Present: __________    Number of Visitors Present: __________    Total: _____\n\n")
        f.write("MEC/VEC/ECs Present:\n")
        f.writelines(["    ____________________________________________\n" for _ in range(3)])
        for index, category in enumerate(CATEGORIES, 1):
            f.write(f"{index}. {category}\n")
            for item in agenda_items[category]:
                f.write(f"    - {item}\n")
            f.write("\n")
        f.write("Closing Time: __________________\n")
    print(f"Minutes created: {filename}")

# --- Main Execution ---

# Choose group
print("Select the group for the meeting:")
print("1. Oriental-Rabboni Chapter No. 78 RAM")
print("2. Jeremiah Council No. 48")
print("3. Webster Groves Lodge No. 84")
print("4. Ivanhoe Commandery No. 8")
print("5. Other (Enter your own group name)")
group_dict = {
    "1": "Oriental-Rabboni Chapter No. 78 RAM",
    "2": "Jeremiah Council No. 48",
    "3": "Webster Groves Lodge No. 84",
    "4": "Ivanhoe Commandery No. 8"
}
group_choice = input("Enter the number of the group: ").strip()
if group_choice == "5":
    group_name = input("Enter your group name: ").strip()
else:
    group_name = group_dict.get(group_choice, "Unknown Group")

# Meeting date
meeting_date_input = input("Enter the date of the meeting (YYYY-MM-DD): ")
formatted_date = datetime.strptime(meeting_date_input, '%Y-%m-%d').strftime('%B %d, %Y')

# Word or Text format
word_choice = input("Use Microsoft Word format? (y/N): ").strip().lower()
file_ext = ".docx" if word_choice == "y" else ".txt"
agenda_filename = f"agenda_{meeting_date_input}{file_ext}"
minutes_filename = f"minutes_{meeting_date_input}{file_ext}"

# Collect agenda items and generate documents
agenda_items = collect_agenda_items()

if file_ext == ".docx":
    create_agenda_docx(agenda_items, agenda_filename, formatted_date, group_name)
    create_minutes_docx(agenda_items, minutes_filename, formatted_date, group_name)
else:
    create_agenda_txt(agenda_items, agenda_filename, formatted_date, group_name)
    create_minutes_txt(agenda_items, minutes_filename, formatted_date, group_name)
